"""
@doc
Core service для генерации DOCX документов из JSON профилей должностей.

Перемещен из services в core как domain service - генерация DOCX
является частью бизнес-логики, не внешним сервисом.

Конвертирует структурированные JSON профили в красиво отформатированные
DOCX документы с таблицами, разделами и корпоративным стилем.

Examples:
  python> service = initialize_docx_service()
  python> docx_path = service.create_docx_from_json(profile_data, "profile.docx")
"""

import os
import logging
from datetime import datetime
from typing import Dict, Any, Optional

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("⚠️ python-docx not installed. DOCX generation will be disabled.")
    DOCX_AVAILABLE = False


class ProfileDocxService:
    """
    @doc
    Core service для генерации DOCX документов из JSON профилей должностей.

    Создает профессионально оформленные DOCX файлы с:
    - Корпоративным стилем
    - Структурированными таблицами
    - Заголовками и разделами
    - Правильным форматированием

    Examples:
      python> service = ProfileDocxService()
      python> path = service.create_docx_from_json(json_data, "output.docx")
    """

    def __init__(self):
        """Инициализация сервиса"""
        self.available = DOCX_AVAILABLE
        if not self.available:
            logger.warning("DOCX service initialized but python-docx is not available")

    def create_docx_from_json(self, json_data: Dict[str, Any], output_path: str) -> str:
        """
        @doc
        Создает DOCX документ из JSON профиля.

        Args:
          json_data: Данные профиля в JSON формате
          output_path: Путь для сохранения DOCX файла

        Returns:
          str: Путь к созданному DOCX файлу

        Examples:
          python> path = service.create_docx_from_json(profile_data, "/tmp/profile.docx")
        """
        if not self.available:
            raise RuntimeError("python-docx is not installed. Cannot generate DOCX files.")

        try:
            # Извлекаем профиль из структуры данных
            if "profile" in json_data:
                profile = json_data.get("profile", {})
            else:
                profile = json_data

            # Создаем новый документ
            doc = Document()

            # Настраиваем стили документа
            self._setup_document_styles(doc)

            # Генерируем содержание документа
            self._add_header(doc, profile)
            self._add_basic_info(doc, profile)
            self._add_responsibilities(doc, profile)
            self._add_skills(doc, profile)
            self._add_personal_qualities(doc, profile)
            self._add_corporate_competencies(doc, profile)
            self._add_education(doc, profile)
            self._add_career_development(doc, profile)
            self._add_workplace_provisioning(doc, profile)
            self._add_performance_metrics(doc, profile)
            self._add_additional_info(doc, profile)
            self._add_metadata(doc, json_data)

            # Сохраняем документ
            doc.save(output_path)
            logger.info(f"✅ DOCX файл создан: {output_path}")

            return output_path

        except Exception as e:
            logger.error(f"❌ Ошибка создания DOCX: {e}")
            raise

    def _setup_document_styles(self, doc: "Document") -> None:
        """Настраивает стили документа"""
        try:
            # Настраиваем основные стили
            styles = doc.styles

            # Стиль заголовка
            if 'Heading 1' in styles:
                heading1 = styles['Heading 1']
                heading1.font.size = Pt(16)
                heading1.font.color.rgb = RGBColor(0, 51, 102)  # Темно-синий

            # Стиль подзаголовка
            if 'Heading 2' in styles:
                heading2 = styles['Heading 2']
                heading2.font.size = Pt(14)
                heading2.font.color.rgb = RGBColor(0, 102, 204)  # Синий

        except Exception as e:
            logger.warning(f"⚠️ Не удалось настроить стили: {e}")

    def _add_header(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет заголовок документа"""
        title = profile.get("position_title", "Должность не указана")
        department = profile.get("department_specific", "")

        # Основной заголовок
        header = doc.add_heading(f"📋 {title}", level=1)
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if department:
            dept_p = doc.add_paragraph(f"Подразделение: {department}")
            dept_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Добавляем разделитель
        doc.add_paragraph("_" * 80)

    def _add_basic_info(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет основную информацию"""
        doc.add_heading("📊 Основная информация", level=2)

        # Создаем таблицу
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Заголовки таблицы
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'

        # Данные для таблицы
        basic_info = [
            ("Название должности", profile.get("position_title", "Не указано")),
            ("Блок", profile.get("department_broad", "Не указано")),
            ("Департамент/Отдел", profile.get("department_specific", "Не указано")),
            ("Категория должности", profile.get("position_category", profile.get("category", "Не указано"))),
            ("Непосредственный руководитель", profile.get("direct_manager", "Не указано")),
            ("Тип деятельности", profile.get("primary_activity_type", profile.get("primary_activity", "Не указано"))),
        ]

        # Подчиненные
        subordinates = profile.get("subordinates", {})
        if isinstance(subordinates, dict):
            departments = subordinates.get("departments", 0)
            people = subordinates.get("direct_reports", subordinates.get("people", 0))
            subordinates_text = f"Департаментов: {departments}, Человек: {people}"
        else:
            subordinates_text = str(subordinates)
        basic_info.append(("Подчиненные", subordinates_text))

        # Заполняем таблицу
        for param, value in basic_info:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)

    def _add_responsibilities(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет области ответственности"""
        doc.add_heading("🎯 Области ответственности", level=2)

        responsibilities = profile.get("responsibility_areas", [])
        if not responsibilities:
            doc.add_paragraph("Области ответственности не определены", style='Italic')
            return

        for i, area in enumerate(responsibilities, 1):
            if isinstance(area, dict):
                area_name = area.get("area")
                if area_name is None:
                    area_name = area.get("title", f"Область {i}")
                elif isinstance(area_name, list):
                    area_name = area_name[0] if area_name else f"Область {i}"

                doc.add_heading(f"{i}. {area_name}", level=3)

                # Задачи
                tasks = area.get("tasks", [])
                if tasks:
                    for task in tasks:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(str(task))

    def _add_skills(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет профессиональные навыки"""
        doc.add_heading("🛠️ Профессиональные навыки", level=2)

        skills = profile.get("professional_skills", [])
        if not skills:
            doc.add_paragraph("Профессиональные навыки не определены", style='Italic')
            return

        for skill_category in skills:
            if isinstance(skill_category, dict):
                category_name = skill_category.get("skill_category", skill_category.get("category", "Неизвестная категория"))
                doc.add_heading(category_name, level=3)

                specific_skills = skill_category.get("specific_skills", skill_category.get("skills", []))

                if specific_skills and len(specific_skills) > 0 and isinstance(specific_skills[0], dict):
                    # Детальные навыки с уровнями - создаем таблицу
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Table Grid'

                    # Заголовки
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Навык'
                    hdr_cells[1].text = 'Уровень'
                    hdr_cells[2].text = 'Описание'

                    for skill in specific_skills:
                        name = skill.get("skill_name", "Неизвестный навык")
                        level = skill.get("proficiency_level", skill.get("target_level", "Не указан"))
                        description = skill.get("proficiency_description", "Описание отсутствует")

                        # Конвертируем числовой уровень в текст
                        if isinstance(level, int):
                            level_map = {1: "Базовый", 2: "Средний", 3: "Продвинутый", 4: "Экспертный"}
                            level_text = level_map.get(level, f"Уровень {level}")
                        else:
                            level_text = str(level)

                        row_cells = table.add_row().cells
                        row_cells[0].text = name
                        row_cells[1].text = level_text
                        row_cells[2].text = description
                else:
                    # Простой список навыков
                    for skill in specific_skills:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(str(skill))

    def _add_personal_qualities(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет личностные качества"""
        doc.add_heading("👤 Личностные качества", level=2)

        qualities = profile.get("personal_qualities", [])
        if not qualities:
            doc.add_paragraph("Личностные качества не определены", style='Italic')
            return

        # Добавляем качества списком
        for quality in qualities:
            p = doc.add_paragraph(style='List Bullet')
            if isinstance(quality, dict):
                quality_name = quality.get("quality", quality.get("name", "Качество"))
                p.add_run(quality_name.capitalize())
            else:
                p.add_run(str(quality).capitalize())

    def _add_corporate_competencies(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет корпоративные компетенции"""
        doc.add_heading("🏢 Корпоративные компетенции", level=2)

        competencies = profile.get("corporate_competencies", [])
        if not competencies:
            doc.add_paragraph("Корпоративные компетенции не определены", style='Italic')
            return

        # Добавляем компетенции списком
        for competency in competencies:
            p = doc.add_paragraph(style='List Bullet')
            if isinstance(competency, dict):
                comp_name = competency.get("competency", competency.get("name", "Компетенция"))
                p.add_run(comp_name)
            else:
                p.add_run(str(competency))

    def _add_education(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет требования к образованию"""
        doc.add_heading("🎓 Образование и опыт работы", level=2)

        education = profile.get("experience_and_education", profile.get("education_requirements", profile.get("education", {})))
        if not education:
            doc.add_paragraph("Требования к образованию не определены", style='Italic')
            return

        # Создаем таблицу для основной информации
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Требование'
        hdr_cells[1].text = 'Описание'

        edu_info = []

        # Собираем информацию
        if "education_level" in education:
            edu_info.append(("Уровень образования", education["education_level"]))

        field_of_study = education.get("field_of_study")
        if field_of_study:
            edu_info.append(("Область обучения", field_of_study))

        total_work_experience = education.get("total_work_experience", education.get("total_experience"))
        if total_work_experience:
            edu_info.append(("Общий опыт работы", total_work_experience))

        # Заполняем таблицу
        for req, desc in edu_info:
            row_cells = table.add_row().cells
            row_cells[0].text = req
            row_cells[1].text = str(desc)

    def _add_career_development(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет карьерограмму"""
        doc.add_heading("📈 Карьерограмма", level=2)

        careerogram = profile.get("careerogram", {})
        career_legacy = profile.get("career_development", profile.get("career_path", {}))
        career_data = {**career_legacy, **careerogram}

        if not career_data:
            doc.add_paragraph("Информация о карьерном развитии не определена", style='Italic')
            return

        # Входные позиции
        source_positions = career_data.get("source_positions", {})
        if source_positions:
            doc.add_heading("🚪 Входные позиции", level=3)

            if isinstance(source_positions, list):
                for pos in source_positions:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(str(pos))
            elif isinstance(source_positions, dict):
                direct_predecessors = source_positions.get("direct_predecessors", [])
                if direct_predecessors:
                    doc.add_paragraph("Прямые предшественники:", style='Intense Quote')
                    for pos in direct_predecessors:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(str(pos))

    def _add_workplace_provisioning(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет обеспечение рабочего места"""
        doc.add_heading("💻 Обеспечение рабочего места", level=2)

        workplace = profile.get("workplace_provisioning", {})
        tech_req_legacy = profile.get("technical_requirements", {})
        provisioning_data = {**tech_req_legacy, **workplace}

        if not provisioning_data:
            doc.add_paragraph("Требования к обеспечению рабочего места не определены", style='Italic')
            return

        # Программное обеспечение
        software_info = provisioning_data.get("software", {})
        if software_info:
            doc.add_heading("📱 Программное обеспечение", level=3)

            standard_package = software_info.get("standard_package", [])
            if standard_package:
                doc.add_paragraph("Стандартный пакет:", style='Intense Quote')
                for sw in standard_package:
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(str(sw))

    def _add_performance_metrics(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет показатели эффективности"""
        doc.add_heading("📊 Показатели эффективности", level=2)

        metrics = profile.get("performance_metrics", {})
        if not metrics:
            doc.add_paragraph("Показатели эффективности не определены", style='Italic')
            return

        # Методология оценки
        methodology = metrics.get("evaluation_methodology")
        if methodology:
            p = doc.add_paragraph()
            p.add_run("Методология оценки: ").bold = True
            p.add_run(str(methodology))

        # Показатели успеха
        success_indicators = metrics.get("success_indicators", [])
        if success_indicators:
            doc.add_heading("🎯 Показатели успеха", level=3)
            for indicator in success_indicators:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(str(indicator))

    def _add_additional_info(self, doc: "Document", profile: Dict[str, Any]) -> None:
        """Добавляет дополнительную информацию"""
        doc.add_heading("ℹ️ Дополнительная информация", level=2)

        additional = profile.get("additional_information", {})
        if not additional:
            doc.add_paragraph("Дополнительная информация отсутствует", style='Italic')
            return

        # Условия работы
        working_conditions = additional.get("working_conditions", {})
        if working_conditions:
            doc.add_heading("🏢 Условия работы", level=3)

            schedule = working_conditions.get("work_schedule")
            if schedule:
                p = doc.add_paragraph()
                p.add_run("График работы: ").bold = True
                p.add_run(str(schedule))

    def _add_metadata(self, doc: "Document", json_data: Dict[str, Any]) -> None:
        """Добавляет метаданные"""
        doc.add_heading("📋 Метаданные", level=2)

        generation_meta = json_data.get("metadata", {})

        # Создаем таблицу для метаданных
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Заголовки
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Параметр'
        hdr_cells[1].text = 'Значение'

        metadata_info = []

        # Информация о генерации
        generation_info = generation_meta.get("generation", {})
        if generation_info:
            if "timestamp" in generation_info:
                metadata_info.append(("Дата генерации", generation_info["timestamp"]))
            if "duration" in generation_info:
                duration = generation_info["duration"]
                metadata_info.append(("Время генерации", f"{duration:.2f} сек"))

        # LLM информация
        llm_info = generation_meta.get("llm", {})
        if llm_info:
            if "model" in llm_info:
                metadata_info.append(("Модель LLM", llm_info["model"]))

        # Заполняем таблицу
        for param, value in metadata_info:
            row_cells = table.add_row().cells
            row_cells[0].text = param
            row_cells[1].text = str(value)

        # Дата создания DOCX
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        row_cells = table.add_row().cells
        row_cells[0].text = "DOCX создан"
        row_cells[1].text = current_time


def initialize_docx_service() -> Optional[ProfileDocxService]:
    """
    @doc
    Инициализирует и возвращает экземпляр DOCX сервиса.

    Returns:
      Optional[ProfileDocxService]: Экземпляр сервиса или None если python-docx недоступен

    Examples:
      python> service = initialize_docx_service()
      python> if service: service.create_docx_from_json(data, "out.docx")
    """
    try:
        service = ProfileDocxService()
        if service.available:
            logger.info("✅ DOCX Service initialized successfully")
            return service
        else:
            logger.warning("⚠️ DOCX Service not available (python-docx not installed)")
            return None
    except Exception as e:
        logger.error(f"❌ Error initializing DOCX service: {e}")
        return None


if __name__ == "__main__":
    print("✅ ProfileDocxService - Convert JSON profiles to beautiful DOCX documents")
    print("📄 Features: Corporate styling, structured tables, professional layout")
    print("🎨 Output: Clean, readable Word documents")